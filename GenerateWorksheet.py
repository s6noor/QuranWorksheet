from bs4 import BeautifulSoup
import re
import docx
from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import requests

# ------------USER INPUT FOR WHICH SURAH THEY WANT THE WORKSHEET FOR ---------------------
chapter_number = int(input("Enter chapter number you want: "))
if chapter_number > 114 or chapter_number < 1:
    raise ValueError("Please enter a number between 1 and 114")

#VARIABLES
url = "https://api.quran.com/api/v4/quran"
VERSES = "/verses/uthmani"
TAFSIR = "/tafsirs/{tafsir_id}" #resources = json.loads(requests.get("https://api.quran.com/api/v4/resources/tafsirs").text)
VERSE_TRANSLATION = "/translations/{id}" #resources = json.loads(requests.get("https://api.quran.com/api/v4/resources/translations").text)

translation_lookup = {'id': 131, 'name': "The Clear Quran"} #using The Clear Quran Translation - future work to make this more modular
tafsir_lookup = {'id': 160, 'name': "Tafsir Ibn Kathir"} #using Tafsir Ibn Kathir - future work to make this modular

#API parameters
payload = {"chapter_number": str(chapter_number)}

#---------------------------------------------------------------------------------------------------------------------------------------#
# Note - I know that I havent done error handling here -- future work

#Get each ayah from requested chapter_number
ayah= json.loads(requests.get(url+VERSES, params=payload).text)
verses_list = ayah['verses']
#print(verses_list)

#Get translation of each ayah from requested chapter_number
translation = json.loads(requests.get(url+VERSE_TRANSLATION.format(id=translation_lookup['id']), params=payload).text)
translation_list = translation['translations']
#print(translation)

if len(verses_list) != len(translation_list):
    raise ValueError("The number of ayah is the verses and translation are not the same. Please check the chapter number, and IDs")

#Request for Tafsir Info that will go on the first Page
payload['language'] = 'english'
tafsir = json.loads(requests.get(url+TAFSIR.format(tafsir_id = tafsir_lookup["id"]), params=payload).text)
tafsir_text = tafsir['tafsirs'][0]["text"] #assuming that first one on the list is an english tafsir
#print(tafsir_text)

#----------------------------------------------------------------------------------------------------------------------------------------#
# Creat worksheet with first presenting a short tafsir, then one ayah per page with its translation. Leave whitespace for student notes

doc = Document()

#--------------------FIRST ADD TAFSIR ----------------------
#tafsir and translation text has html headers, so using a parser here to present text properly
soup = BeautifulSoup(tafsir_text, 'html.parser')

# Extract and add the HTML content to the document
# Code modified from ChatGPT
for tag in soup.find_all(['h1', 'h2', 'h3', 'p']):
    if tag.name == 'h1':
        heading = doc.add_heading(tag.text, level=1)
        heading.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif tag.name == 'h2':
        doc.add_heading(tag.text, level=2)
    elif tag.name == 'h3':
        doc.add_heading(tag.text, level=3)
    elif tag.name == 'p':
        paragraph = doc.add_paragraph(tag.text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


# ------------------ ADD AYAT AND TRANSLATION ON EACH PAGE --------------------------
# Structure - the verse/translation lists are a list of dictionaries, with each dictionary corresponding to one ayah.
for verse, translation in zip(verses_list, translation_list):
    
    #add each ayah into a new page
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENTATION.LANDSCAPE #this doesnt work
    
    
    arabic_text = verse['text_uthmani']
    translation_text = translation['text'] 
    translation_text = re.sub(r'<sup.*?</sup>', '', translation_text) #remove this if you want the foot note signs
        
    # Add the Ayat as a paragraph to the document
    # Note that if you dont have an appropriate arabic font downloaded, you might have trouble. I used this to download Al_Mushaf: https://urdunigaar.com/download-quranic-font-quran-standard-font-islamic-fonts/ 
    paragraph = doc.add_paragraph()
    paragraph.add_run("\n\n\n")
    run = paragraph.add_run(arabic_text)
    run.font.size = docx.shared.Pt(32)
    paragraph.add_run("\n\n\n")
    
    #Add translation under the ayah after appropriate spacing. 
    run = paragraph.add_run(translation_text)
    run.font.size = docx.shared.Pt(20)
    paragraph.add_run("\n")
    run = paragraph.add_run(translation_lookup['name']) #also adding translation name
    run.font.size = docx.shared.Pt(15)
    
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT #align everything to the right

# Save the worksheet based on Surah Nummber 
file_name = "Surah " + str(chapter_number) + " Worksheet.docx"
doc.save(file_name)